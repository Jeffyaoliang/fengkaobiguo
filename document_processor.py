import os
import PyPDF2
from docx import Document
import markdown
from bs4 import BeautifulSoup
from typing import List, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def read_text_file(self, file_path: str) -> str:
        """读取文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"读取文本文件失败: {e}")
            return ""
    
    def read_pdf_file(self, file_path: str) -> str:
        """读取PDF文件"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"读取PDF文件失败: {e}")
            return ""
    
    def read_docx_file(self, file_path: str) -> str:
        """读取Word文档"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"读取Word文档失败: {e}")
            return ""
    
    def read_markdown_file(self, file_path: str) -> str:
        """读取Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
                # 转换为纯文本
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
        except Exception as e:
            logger.error(f"读取Markdown文件失败: {e}")
            return ""
    
    def process_file(self, file_path: str) -> List[LangchainDocument]:
        """处理单个文件并返回文档块"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # 根据文件类型选择读取方法
        if file_extension == '.txt':
            content = self.read_text_file(file_path)
        elif file_extension == '.pdf':
            content = self.read_pdf_file(file_path)
        elif file_extension == '.docx':
            content = self.read_docx_file(file_path)
        elif file_extension == '.md':
            content = self.read_markdown_file(file_path)
        else:
            logger.warning(f"不支持的文件格式: {file_extension}")
            return []
        
        if not content.strip():
            logger.warning(f"文件内容为空: {file_path}")
            return []
        
        # 创建Langchain文档对象
        doc = LangchainDocument(
            page_content=content,
            metadata={
                "source": file_path,
                "file_type": file_extension,
                "file_name": os.path.basename(file_path)
            }
        )
        
        # 分割文档
        chunks = self.text_splitter.split_documents([doc])
        logger.info(f"文件 {file_path} 被分割为 {len(chunks)} 个块")
        
        return chunks
    
    def process_directory(self, directory_path: str) -> List[LangchainDocument]:
        """处理目录中的所有支持的文件"""
        all_chunks = []
        supported_formats = ['.txt', '.pdf', '.docx', '.md']
        
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_extension = os.path.splitext(file)[1].lower()
                
                if file_extension in supported_formats:
                    logger.info(f"处理文件: {file_path}")
                    chunks = self.process_file(file_path)
                    all_chunks.extend(chunks)
        
        logger.info(f"总共处理了 {len(all_chunks)} 个文档块")
        return all_chunks 